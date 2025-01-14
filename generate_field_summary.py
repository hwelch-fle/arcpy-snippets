import arcpy
import arcpy.typing
import arcpy.typing.describe
import docx
from datetime import datetime
from pathlib import Path

# Utility functions
def docx_find_replace_text(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old_text in paragraph.text:
                        for run in paragraph.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)

def replace_date_in_footer(doc):
    current_date = datetime.now().strftime("%d %B %Y")
    for section in doc.sections:
        footer = section.footer
        for paragraph in footer.paragraphs:
            if "<<DATE2>>" in paragraph.text:
                for run in paragraph.runs:
                    if "<<DATE2>>" in run.text:
                        run.text = run.text.replace("<<DATE2>>", current_date)

def as_dict(cursor: arcpy.da.SearchCursor):
    """Convert a search cursor to a dictionary."""
    yield from (dict(zip(cursor.fields, row)) for row in cursor)
    
def get_domains(feature_class) -> dict[str, arcpy.da.Domain]:
    """Return a dictionary of domains in the workspace."""
    fc_desc = arcpy.Describe(feature_class)
    return {domain.name: domain for domain in arcpy.da.ListDomains(fc_desc.workspace.catalogPath)}

def generate_field_summary(inputfilename, outputfolder, feature_class):
    # Define paths
    inputfilename = Path(inputfilename)
    outputfolder = Path(outputfolder)
    feature_class = Path(feature_class)
    
    # Define field mapping
    field_mapping = {
        "<<PROJECT_NAME>>": "PROJECT_NAME",
        "<<FIELD_DATE>>": "FIELD_DATE",
        "<<ARCH_CREW>>": "ARCH_CREW",
        "<<PERMIT>>": "PERMIT",
        "<<DIVISION>>": "DIVISION",
        "<<METHOD>>": "METHOD",
        "<<DIST_EXIST>>": "DIST_EXIST",
        "<<DESCRIPTION>>": "DESCRIPTION",
        "<<DIST_REQ>>": "DIST_REQ",
        "<<HISTORY>>": "HISTORY",
        "<<SUB_OB>>": "SUB_OB",
        "<<ARCH_OB>>": "ARCH_OB",
        "<<REC>>": "REC"
    }

    # Subtype to domain mapping
    subtype_to_domain = {
        0: "DBLBranch",
        1: "ParksBranch",
        2: "ParksBranch",
        3: "REFMBranch",
        4: "CMOBranch"
    }

    subtype_field = "DEPT"  # Subtype field
    feature_fields = list(field_mapping.values()) + ['last_edited_date', subtype_field]

    features: list[dict[str, str | int | float]] = [
        row 
        for row in as_dict(arcpy.da.SearchCursor(feature_class, feature_fields))
    ]
    domains = get_domains(feature_class)

    for feature in features:
        last_edited_date = feature['last_edited_date']
        subtype_code = feature[subtype_field]

        if not last_edited_date:
            print("Skipping record with no last_edited_date.")
            continue

        output_file = outputfolder / f"{feature['PROJECT_NAME']}_FieldSummary.docx"
        if output_file.exists():
            doc_mod_time = datetime.fromtimestamp(output_file.stat().st_mtime)
            if doc_mod_time >= last_edited_date:
                print(f"Skipping {output_file}, already up-to-date.")
                continue

        doc = docx.Document(str(inputfilename))
        for placeholder, field in field_mapping.items():
            value = feature[field]

            # Handle DIVISION with subtype-specific domains
            if field != "DIVISION" and value in subtype_to_domain:
                value = domains[subtype_to_domain[subtype_code]]

            docx_find_replace_text(doc, placeholder, str(value))
            # Replace Header
            # docx_replace_header(doc, placeholder, value) # This is not implemented
            
        # Replace date in the footer
        replace_date_in_footer(doc)
        
        doc.save(str(output_file))
        
        print(f"Created or updated: {output_file.name}")
    print("Process completed.")

if __name__ == "__main__":
    from argparse import ArgumentParser
    
    parser = ArgumentParser(
        name="Generate Field Summary",
        description="Generates and maintains field summary documents for fieldwork records."
    )
    parser.add_argument("-i", "--inputfile", help="Input filename")
    parser.add_argument("-o", "--outputfile", help="Output folder")
    parser.add_argument("-f", "-featureclass", help="Feature class")
    args = parser.parse_args()
    
    generate_field_summary(args.inputfile, args.outputfile, args.featureclass)